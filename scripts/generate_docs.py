"""
Generates the project architecture documentation as a Word (.docx) file.
Run: python scripts/generate_docs.py
Output: docs/ARCHITECTURE.docx
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parents[1]))

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run()
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
        run.font.size = Pt(14)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run.font.size = Pt(12)
    return h


def add_para(doc, text, bold=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_code(doc, text):
    """Add a monospaced code block (grey background)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            # alternate row shading
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E9EFF7')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


# ── document ─────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Cover ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Real-Time Financial Risk Scoring Platform")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("Architecture & Technical Documentation")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Dataset: PaySim  |  Model: LightGBM + SHAP  |  Stack: FastAPI · Kafka · Redis · Evidently").font.size = Pt(10)

    doc.add_page_break()

    # ── 1. Project Overview ────────────────────────────────────────────────
    set_heading(doc, "1. Project Overview", 1)
    add_para(doc,
        "This platform scores financial transactions for fraud risk in real time (<50 ms per "
        "transaction). It ingests a stream of payment events, computes features, runs a "
        "LightGBM model, applies business decision rules, and returns an APPROVE / REVIEW / BLOCK "
        "decision together with a SHAP-based explanation of why the transaction was scored that way.")
    doc.add_paragraph()
    add_para(doc, "Core Problem", bold=True)
    add_para(doc,
        "Fraud detection is an extreme class-imbalance problem — 1 fraud per 774 legitimate "
        "transactions (0.13%). Standard accuracy metrics are meaningless. The entire system is "
        "designed around Precision-Recall AUC (PR-AUC).")
    doc.add_paragraph()
    add_para(doc, "Baseline vs Final Model", bold=True)
    add_table(doc,
        ["Model", "PR-AUC", "Notes"],
        [
            ["Logistic Regression", "0.827", "Starting point from EDA notebook"],
            ["LightGBM", "~0.930", "+12% improvement, fully explainable"],
        ],
        col_widths=[2.2, 1.2, 3.6]
    )

    doc.add_paragraph()

    # ── 2. Tech Stack ──────────────────────────────────────────────────────
    set_heading(doc, "2. Tech Stack", 1)
    add_table(doc,
        ["Layer", "Technology", "Why chosen"],
        [
            ["ML Model",           "LightGBM",              "Fast training, native imbalance handling, exact SHAP"],
            ["Explainability",     "SHAP TreeExplainer",    "Exact attribution per prediction in <5 ms"],
            ["Inference API",      "FastAPI + Uvicorn",     "Async, Pydantic validation, auto OpenAPI docs"],
            ["Schema validation",  "Pydantic v2",           "Runtime type checking at every boundary"],
            ["Message broker",     "Apache Kafka",          "Durable, ordered, replayable event stream"],
            ["Online feature store","Redis",                "Sub-millisecond velocity counter lookups"],
            ["Drift detection",    "Evidently",             "Production-grade data & model drift reports"],
            ["Containerisation",   "Docker + Compose",      "Reproducible local and production environment"],
            ["Configuration",      "Environment variables", "12-Factor App — no hardcoded secrets"],
        ],
        col_widths=[1.8, 1.9, 3.3]
    )
    doc.add_paragraph()

    # ── 3. System Architecture ─────────────────────────────────────────────
    set_heading(doc, "3. System Architecture", 1)

    set_heading(doc, "3.1  Three Main Pipelines", 2)
    add_para(doc,
        "The platform is split into three independent pipelines that share only "
        "the data contract (Pydantic schemas) and the model artifact file.")
    doc.add_paragraph()

    add_table(doc,
        ["Pipeline", "Trigger", "Key files", "Latency target"],
        [
            ["Sync (HTTP)",   "REST POST /v1/score",      "api/main.py, feature_pipeline.py, lgbm_model.py",         "<50 ms p99"],
            ["Async (Kafka)", "Message on 'transactions'","kafka_consumer.py, feature_pipeline.py, lgbm_model.py",   "<15 ms/msg"],
            ["Offline",       "Manual / scheduled",       "train.py, retrain_pipeline.py, drift_detector.py",        "Minutes"],
        ],
        col_widths=[1.3, 2.0, 3.2, 1.5]
    )
    doc.add_paragraph()

    set_heading(doc, "3.2  Component Map", 2)
    add_code(doc,
"Real-Time-Financial-Risk/\n"
"├── src/\n"
"│   ├── datasets/\n"
"│   │   ├── schema.py                ← DATA CONTRACTS (Pydantic)\n"
"│   │   └── adapters/\n"
"│   │       └── paysim_adapter.py    ← BATCH LOAD + STREAM REPLAY\n"
"│   ├── features/\n"
"│   │   └── feature_pipeline.py     ← FEATURE ENGINEERING (batch + online)\n"
"│   ├── models/\n"
"│   │   └── lgbm_model.py           ← LIGHTGBM + SHAP + VERSIONING\n"
"│   ├── training/\n"
"│   │   ├── train.py                ← OFFLINE TRAINING SCRIPT\n"
"│   │   └── retrain_pipeline.py     ← AUTOMATED RETRAINING\n"
"│   ├── ingestion/\n"
"│   │   ├── kafka_producer.py       ← TRANSACTION STREAM SIMULATOR\n"
"│   │   └── kafka_consumer.py       ← REAL-TIME SCORING CONSUMER\n"
"│   └── monitoring/\n"
"│       └── drift_detector.py       ← DRIFT DETECTION + SCORE MONITOR\n"
"├── api/\n"
"│   ├── main.py                     ← FASTAPI INFERENCE SERVICE\n"
"│   └── schemas.py                  ← API REQUEST / RESPONSE MODELS\n"
"├── configs/\n"
"│   └── settings.py                 ← ENV-VAR CONFIGURATION\n"
"└── deployment/\n"
"    ├── Dockerfile                  ← MULTI-STAGE PRODUCTION IMAGE\n"
"    ├── docker-compose.yml          ← FULL LOCAL STACK\n"
"    └── requirements.txt            ← ALL DEPENDENCIES"
    )
    doc.add_paragraph()

    # ── 4. Data Flow ──────────────────────────────────────────────────────
    set_heading(doc, "4. Data Flow Diagrams", 1)

    set_heading(doc, "4.1  Sync Path — HTTP API", 2)
    add_code(doc,
"Client                 FastAPI               Redis          LightGBM\n"
"  |                       |                    |                |\n"
"  |-- POST /v1/score ----->|                   |                |\n"
"  |                       |-- GET vel:count -->|                |\n"
"  |                       |<-- count=3 --------|                |\n"
"  |                       |-- GET vel:amount ->|                |\n"
"  |                       |<-- amount=500 -----|                |\n"
"  |                       |-- predict_proba(X) --------------->|\n"
"  |                       |<-- 0.9312 -------------------------|\n"
"  |                       |-- explain(X, top_k=5) ------------>|\n"
"  |                       |<-- {feature: shap_val} -----------|\n"
"  |                       |-- INCR vel:count -->|              |\n"
"  |<-- {score, BLOCK} ----|                    |                |"
    )
    doc.add_paragraph()

    set_heading(doc, "4.2  Async Path — Kafka Stream", 2)
    add_code(doc,
"PaySim Adapter      Kafka Broker        Consumer Group       Redis\n"
"      |                  |                    |                 |\n"
"      |-- publish(tx) -->|                    |                 |\n"
"      |-- publish(tx) -->|                    |                 |\n"
"      |                  |-- poll(100 msgs) -->|                |\n"
"      |                  |                    |-- GET velocity ->|\n"
"      |                  |                    |<-- counts -------|\n"
"      |                  |                    |                 |\n"
"      |                  |                    | score all 100   |\n"
"      |                  |<-- publish results-|                 |\n"
"      |                  |  (fraud-decisions) |                 |\n"
"      |                  |-- ack offset ----->|                 |"
    )
    doc.add_paragraph()

    set_heading(doc, "4.3  Training Pipeline Flow", 2)
    add_code(doc,
"PaySim CSV (6.36M rows)\n"
"         |\n"
"   batch_load()  -- time-ordered split at step=500\n"
"         |\n"
"    +----+---------------------+\n"
"    |                          |\n"
" train_df                   test_df\n"
" steps 1-500                steps 501-743\n"
" 6.06M rows / 5,561 frauds  300K rows / 2,652 frauds\n"
"    |\n"
" engineer_features_batch()   <- 10 features computed\n"
"    |\n"
" LightGBM.fit(\n"
"   scale_pos_weight = 774    <- handles 0.13% fraud rate\n"
"   early_stopping = 50       <- stops when val PR-AUC plateaus\n"
" )\n"
"    |\n"
" Build SHAP TreeExplainer    <- exact attributions, ~3ms\n"
"    |\n"
" Evaluate on test set\n"
"   PR-AUC: ~0.93\n"
"   Precision @ 0.5: ~0.99\n"
"   Recall    @ 0.5: ~0.75\n"
"    |\n"
" model.save('v1')\n"
"   -> src/models/artifacts/fraud_model_v1.joblib\n"
"   -> src/models/registry/champion.json"
    )
    doc.add_paragraph()

    set_heading(doc, "4.4  Monitoring & Retraining Loop", 2)
    add_code(doc,
"'fraud-decisions' Kafka topic\n"
"         |\n"
"   ScoreMonitor.record()       <- rolling window of 10,000 decisions\n"
"         |\n"
"         +-- every 1,000: check block_rate, latency_p99 -> alert if anomalous\n"
"         |\n"
"         +-- every 50,000: DriftDetector.check_drift()\n"
"                              compare feature distributions\n"
"                              (Evidently KS test per feature)\n"
"                                    |\n"
"                             drift_ratio > 0.3?\n"
"                                    |\n"
"                                   YES\n"
"                                    |\n"
"                         retrain_pipeline.run(trigger='drift')\n"
"                                    |\n"
"                           new PR-AUC > champion + 0.005?\n"
"                                    |\n"
"                          +----YES--+--NO----+\n"
"                          |                  |\n"
"                       promote             discard\n"
"                       to champion         challenger\n"
"                          |\n"
"                   API restart picks up\n"
"                   new MODEL_VERSION"
    )
    doc.add_paragraph()

    # ── 5. Component Breakdown ────────────────────────────────────────────
    set_heading(doc, "5. Component-by-Component Breakdown", 1)

    components = [
        (
            "5.1  Data Contracts — src/datasets/schema.py",
            "Defines three Pydantic models that act as the shared contract between every layer. "
            "Nothing in the system passes raw dicts — everything is validated at runtime.",
            [
                ("TransactionEvent",    "Raw input exactly as it arrives from Kafka or the API. Mirrors PaySim CSV columns."),
                ("EnrichedTransaction", "Transaction after feature engineering — the 10 features the model actually sees."),
                ("RiskScore",           "Final output: score, decision, SHAP top features, latency. Includes a validator that checks decision is consistent with score."),
            ]
        ),
        (
            "5.2  PaySim Adapter — src/datasets/adapters/paysim_adapter.py",
            "Bridges the static CSV file to the streaming interface used by the rest of the system.",
            [
                ("batch_load()",       "Returns (train_df, test_df) split at step=500 for offline training."),
                ("stream(split)",      "Yields TransactionEvent objects one-by-one — identical interface to a Kafka consumer."),
                ("fraud_only_stream()","Yields only fraudulent transactions for stress-testing the decision engine."),
            ]
        ),
        (
            "5.3  Feature Engineering — src/features/feature_pipeline.py",
            "The most critical file. Produces identical feature vectors in both training (batch pandas) and "
            "inference (single event + Redis) modes. This prevents training-serving skew.",
            [
                ("error_balance_orig",  "(oldBal - newBal) - amount. Non-zero when the originator's books don't add up."),
                ("error_balance_dest",  "(newBal - oldBal) - amount. Non-zero when destination receives unexpected amount."),
                ("type_* dummies",      "One-hot encoded transaction type. Fraud only occurs in TRANSFER and CASH_OUT."),
                ("orig_tx_count_1h",    "How many transactions this account sent in the last hour (Redis INCR)."),
                ("orig_amount_sum_1h",  "Total amount sent this hour — detects high-value velocity."),
                ("dest_tx_count_1h",    "How many transfers this destination received — detects mule accounts."),
            ]
        ),
        (
            "5.4  LightGBM Model — src/models/lgbm_model.py",
            "Wraps LightGBM with SHAP explanations and versioned persistence.",
            [
                ("scale_pos_weight=774",   "Computed from actual class ratio. Tells the model missing a fraud costs 774x more."),
                ("min_child_samples=50",   "No leaf can be created from <50 samples — the key regulariser for imbalanced data."),
                ("TreeExplainer",          "Exact SHAP values from the tree structure. Runs in ~3ms. Used for regulatory explainability."),
                ("save/load(version)",     "Artifact stored as fraud_model_v1.joblib. Rollback = change MODEL_VERSION env var."),
            ]
        ),
        (
            "5.5  Inference API — api/main.py",
            "FastAPI service. Model loaded once at startup and reused across all requests. "
            "Decision thresholds are environment variables, not hardcoded.",
            [
                ("POST /v1/score",       "Validate → features → LightGBM → SHAP → APPROVE/REVIEW/BLOCK. <50ms target."),
                ("POST /v1/score/batch", "Score up to 1,000 transactions, sharing SHAP explainer overhead."),
                ("GET /health",          "Liveness probe used by Kubernetes and load balancers."),
            ]
        ),
        (
            "5.6  Retraining Pipeline — src/training/retrain_pipeline.py",
            "Champion/challenger model lifecycle. Only promotes a new model when it demonstrably improves.",
            [
                ("Sliding window",       "Train only on the most recent N steps — recent fraud patterns matter more than old ones."),
                ("Promotion threshold",  "New model must improve PR-AUC by ≥0.5% over the champion to be deployed."),
                ("Registry",             "Every training run logged to src/models/registry/. champion.json = currently live version."),
            ]
        ),
        (
            "5.7  Drift Detection — src/monitoring/drift_detector.py",
            "Two monitors running on the stream of decisions.",
            [
                ("DriftDetector",  "KS test per feature (via Evidently). Flags drift if >30% of features shift significantly."),
                ("ScoreMonitor",   "Rolling window of block_rate and latency_p99. Alerts on 10x block rate surge or >100ms p99."),
            ]
        ),
    ]

    for title, description, items in components:
        set_heading(doc, title, 2)
        add_para(doc, description)
        doc.add_paragraph()
        for name, detail in items:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
            run_name = p.add_run(f"{name}  ")
            run_name.bold = True
            run_name.font.size = Pt(10.5)
            run_detail = p.add_run(detail)
            run_detail.font.size = Pt(10.5)
        doc.add_paragraph()

    # ── 6. Feature Engineering Deep Dive ──────────────────────────────────
    set_heading(doc, "6. Feature Engineering Deep Dive", 1)

    set_heading(doc, "6.1  Why error_balance_orig is the strongest signal", 2)
    add_code(doc,
"LEGITIMATE transaction (TRANSFER $100):\n"
"  oldbalanceOrg  = 500\n"
"  newbalanceOrig = 400   <- decreased by exactly $100\n"
"  amount         = 100\n"
"  error_balance_orig = (500 - 400) - 100 = 0   <- books balance\n\n"
"FRAUDULENT transaction (fabricated amount):\n"
"  oldbalanceOrg  = 181\n"
"  newbalanceOrig = 0     <- drained completely\n"
"  amount         = 500   <- claimed more than was available\n"
"  error_balance_orig = (181 - 0) - 500 = -319  <- STRONGLY non-zero"
    )
    doc.add_paragraph()
    add_para(doc,
        "In the EDA: ~50% of fraud transactions show newbalanceDest = 0 — destination balance "
        "never increases despite supposedly receiving funds. This creates a systematic pattern "
        "that tree models learn as a single split condition.")
    doc.add_paragraph()

    set_heading(doc, "6.2  Velocity Features and Redis", 2)
    add_code(doc,
"In-memory (training / testing):\n"
"  _velocity_store = {\n"
"    'C1231006815': {'step': 501, 'count': 3, 'amount_sum': 900.0}\n"
"  }\n\n"
"Redis (production):\n"
"  key: 'vel:count:C1231006815:501'  -> 3\n"
"  key: 'vel:amount:C1231006815:501' -> 900.0\n"
"  TTL: 7200 seconds  (auto-expires, no manual cleanup needed)\n\n"
"Why Redis and not in-memory in production?\n"
"  - Multiple consumer instances each have separate in-memory state\n"
"  - Redis is shared across all instances -> consistent velocity\n"
"  - LRU eviction handles memory pressure automatically\n"
"  - Atomic INCR prevents race conditions under concurrent writes"
    )
    doc.add_paragraph()

    # ── 7. Model Design Decisions ─────────────────────────────────────────
    set_heading(doc, "7. Model Design Decisions", 1)

    set_heading(doc, "7.1  Why PR-AUC instead of ROC-AUC", 2)
    add_table(doc,
        ["Metric", "Behaviour on 0.13% fraud rate", "Verdict"],
        [
            ["Accuracy",  "'Predict all legitimate' gives 99.87% accuracy",              "Useless"],
            ["ROC-AUC",   "Inflated by massive true-negative count. Easy to get 0.98+",  "Misleading"],
            ["PR-AUC",    "Only counts the fraud class. 0.93 = genuinely hard to achieve","Use this"],
        ],
        col_widths=[1.5, 3.5, 1.5]
    )
    doc.add_paragraph()

    set_heading(doc, "7.2  Why SHAP over alternatives", 2)
    add_table(doc,
        ["Approach", "Speed", "Exact?", "Per-prediction?", "Verdict"],
        [
            ["Feature importance (gain)", "<1 ms", "No",  "No",  "Global only, useless for decisions"],
            ["LIME",                       "~50 ms","No",  "Yes", "Too slow, approximate"],
            ["SHAP TreeExplainer",         "~3 ms", "Yes", "Yes", "Use this — satisfies CFPB/FCA"],
        ],
        col_widths=[2.2, 0.9, 0.9, 1.4, 2.6]
    )
    doc.add_paragraph()

    set_heading(doc, "7.3  Decision Thresholds", 2)
    add_code(doc,
"score  0.0 ─────────── 0.3 ─────────── 0.7 ──── 1.0\n"
"             APPROVE        REVIEW           BLOCK\n\n"
"Controlled by environment variables:\n"
"  REVIEW_THRESHOLD = 0.3   (default)\n"
"  BLOCK_THRESHOLD  = 0.7   (default)\n\n"
"These are business decisions, not ML decisions:\n"
"  'What does it cost to block a good customer?'  (false positive)\n"
"  'What does it cost to let fraud through?'      (false negative)\n"
"  Adjust thresholds based on your risk appetite, not model metrics."
    )
    doc.add_paragraph()

    # ── 8. API Reference ─────────────────────────────────────────────────
    set_heading(doc, "8. API Reference", 1)

    set_heading(doc, "8.1  POST /v1/score — Request", 2)
    add_code(doc,
'{\n'
'  "step":          501,\n'
'  "type":          "TRANSFER",\n'
'  "amount":        181.00,\n'
'  "nameOrig":      "C1231006815",\n'
'  "oldbalanceOrg": 181.00,\n'
'  "newbalanceOrig": 0.00,\n'
'  "nameDest":      "C1666544295",\n'
'  "oldbalanceDest": 0.00,\n'
'  "newbalanceDest": 0.00\n'
'}'
    )
    doc.add_paragraph()

    set_heading(doc, "8.2  POST /v1/score — Response", 2)
    add_code(doc,
'{\n'
'  "step":        501,\n'
'  "name_orig":   "C1231006815",\n'
'  "name_dest":   "C1666544295",\n'
'  "amount":      181.0,\n'
'  "type":        "TRANSFER",\n'
'  "risk_score":  0.9312,\n'
'  "decision":    "BLOCK",\n'
'  "top_features": {\n'
'    "error_balance_orig": 0.4231,    <- strongest fraud signal\n'
'    "type_TRANSFER":      0.3102,    <- transaction type\n'
'    "amount":             0.1823,    <- transaction size\n'
'    "error_balance_dest": -0.0941,   <- negative = pushed toward legit\n'
'    "orig_tx_count_1h":   0.0312\n'
'  },\n'
'  "latency_ms": 6.4\n'
'}'
    )
    doc.add_paragraph()

    set_heading(doc, "8.3  All Endpoints", 2)
    add_table(doc,
        ["Method", "Path", "Description", "Auth"],
        [
            ["GET",  "/health",         "Liveness probe — returns model_loaded status",          "None"],
            ["POST", "/v1/score",       "Score a single transaction, returns SHAP explanation",  "None"],
            ["POST", "/v1/score/batch", "Score up to 1,000 transactions in one call",           "None"],
        ],
        col_widths=[0.8, 1.8, 3.5, 0.9]
    )
    doc.add_paragraph()

    # ── 9. How to Run ─────────────────────────────────────────────────────
    set_heading(doc, "9. How to Run — Step by Step", 1)

    steps = [
        ("Step 1 — Install dependencies",
         "source venv/bin/activate\npip install -r deployment/requirements.txt"),
        ("Step 2 — Train the model  (~3 minutes)",
         "python -m src.training.train --version v1\n# Output: src/models/artifacts/fraud_model_v1.joblib\n# Expected PR-AUC: ~0.93"),
        ("Step 3 — Verify end-to-end (no Docker needed)",
         "python scripts/run_pipeline.py\n# Runs: load -> features -> train -> score 500 txns -> drift check"),
        ("Step 4 — Start the API",
         "uvicorn api.main:app --reload --port 8000\n# Swagger UI: http://localhost:8000/docs"),
        ("Step 5 — Score a transaction",
         'curl -X POST http://localhost:8000/v1/score \\\n'
         '  -H "Content-Type: application/json" \\\n'
         '  -d \'{"step":501,"type":"TRANSFER","amount":181.0,\n'
         '       "nameOrig":"C1231006815","oldbalanceOrg":181.0,"newbalanceOrig":0.0,\n'
         '       "nameDest":"C1666544295","oldbalanceDest":0.0,"newbalanceDest":0.0}\''),
        ("Step 6 — Full streaming stack",
         "docker-compose -f deployment/docker-compose.yml up -d kafka zookeeper redis\n"
         "python -m src.ingestion.kafka_producer --split test --max-events 5000\n"
         "python -m src.ingestion.kafka_consumer --redis-url redis://localhost:6379/0"),
        ("Step 7 — Trigger retraining",
         "python -m src.training.retrain_pipeline --trigger manual\n"
         "# Only promotes if new model PR-AUC improves by >= 0.5%"),
    ]

    for title, code in steps:
        set_heading(doc, title, 2)
        add_code(doc, code)
        doc.add_paragraph()

    # ── 10. What's Not Built Yet ──────────────────────────────────────────
    set_heading(doc, "10. What's Not Built Yet", 1)
    add_table(doc,
        ["Component", "Description", "Priority"],
        [
            ["Graph fraud ring detection", "src/graph/ is empty. Node2Vec / GraphSAGE to detect money-mule networks by analysing the transaction graph.", "High"],
            ["Anomaly detection",          "Isolation Forest for zero-day fraud patterns not present in training data.",                                    "Medium"],
            ["Prometheus metrics",         "/metrics endpoint with request counters, latency histograms, block rate gauges.",                              "Medium"],
            ["Airflow DAGs",               "Scheduled trigger for retrain_pipeline.py. Currently triggered manually only.",                                "Medium"],
            ["Feast feature store",        "Feature versioning and point-in-time correct joins. Currently only Redis is used.",                            "Low"],
            ["Unit tests",                 "Zero test coverage — tests/ directory is empty.",                                                              "High"],
            ["IEEE adapter",               "src/datasets/adapters/ieee_adapter.py is still a stub.",                                                      "Low"],
            ["A/B testing / shadow mode",  "Run two model versions in parallel before promoting, measuring live performance.",                             "Medium"],
        ],
        col_widths=[2.0, 4.0, 1.0]
    )
    doc.add_paragraph()

    # ── Save ──────────────────────────────────────────────────────────────
    out = Path(__file__).parents[1] / "docs" / "ARCHITECTURE.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved: {out}")
    return out


if __name__ == "__main__":
    build_doc()
